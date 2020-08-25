# -----------------------------------------------------------
# ms_word
#
# Defines utility functions to use when writing to assessor agreement template & client contract
# Extends functionality of external docx module
# 
# -----------------------------------------------------------


# External modules

import docx
from docx.shared import Pt


# -----------------------------------------------------------
# General purpose utility functions 
# -----------------------------------------------------------

# Returns docx Document object for specified file

def set_up_doc_object(filepath):
    return docx.Document(filepath)


# Returns docx table object. Table at index number selected from Document.

def set_up_table(doc_obj, table_index):
    return doc_obj.tables[table_index]


# Writes value to specified cell in table 
              
def update_table_cell(table, row, col, value):
    try:
        table.cell(row, col).text = value
    except Exception as e:
        print(table + ' cell' + row + ',' + col + ' not updated with ' + value)
        print(str(e))


# Bolds font in specified cell in table 
# (Cells in reg orgs & sales tracker sheets should only ever have 1 paragraph & 1 run)
                    
def bold_cell(table, row, col):
    try: 
        table.cell(row, col).paragraphs[0].runs[0].font.bold = True
    except Exception as e:
        print(table + ' cell ' + row + ',' + col + ' not formatted bold')
        print(str(e))


# Saves docx Document object (removes .docx file extension & appends org_name + .docx)

def save_document(doc, filepath, org_name, log_message):

    filepath = filepath[:-5]
    filepath = filepath + ' - ' + org_name + '.docx'
    doc.save(filepath)
    print(log_message)


# -----------------------------------------------------------
# Assessor agreement utility functions 
# -----------------------------------------------------------

# Finds docx paragraph with text_to_find and appends provided text as a run.
# Applies correct font to appended run as docx defaults to different font.

def append_text(document, text_to_find, text_to_append):

    try:
        for i in range(len(document.paragraphs)):
            if text_to_find in document.paragraphs[i].text:
                para_to_work_with = document.paragraphs[i]
                para_to_work_with.add_run(' ' + text_to_append)
                appended_run = para_to_work_with.runs[len(para_to_work_with.runs) - 1]
                font = appended_run.font
                font.name = 'Arial'
                font.size = Pt(11)
                font.bold = True
                
    except Exception as e:
        print('Hasn\'t updated ' + text_to_find + ' with ' + text_to_append + ' in Assessor Agreement')
        print(str(e))


# -----------------------------------------------------------
# Client contract utility functions 
# -----------------------------------------------------------

# Finds and replaces specified text with replacement text. Re-formats font to correct font

def find_replace(file, text_to_find, replacement_text):
    
    try:
        for paragraph in file.paragraphs:
            if text_to_find in paragraph.text:
                paragraph.text = paragraph.text.replace(text_to_find, replacement_text)
                for run in paragraph.runs:
                    font = run.font
                    font.name = 'Arial'
                    font.size = Pt(11)
                    font.italic = False
    except Exception as e:               
        print(text_to_find + ' - not replaced with ' + replacement_text + ' in client contract')
        print(str(e))
        

# Sets cell font to Arial 12 Bold 
            
def correct_table_font(table, row, col, bold_boolean):

    try: 
        font = table.cell(row, col).paragraphs[0].runs[0].font
        font.name = 'Arial'
        font.size = Pt(12)
        font.bold = bold_boolean
    except Exception as e:
        print(table + ' cell ' + row + ',' + col + ' font hasn\'t been corrected in client contract')
        print(str(e))
    

# -----------------------------------------------------------
# End
# -----------------------------------------------------------

